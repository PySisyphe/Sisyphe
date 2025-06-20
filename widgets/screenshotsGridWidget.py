"""
External packages/modules
-------------------------

    - PyQt5, Qt GUI, https://www.riverbankcomputing.com/software/pyqt/
    - python-docx, Word document management, https://python-docx.readthedocs.io/en/latest/index.html
"""

from os import getcwd
from os import chdir
from os import remove

from os.path import join
from os.path import exists
from os.path import basename
from os.path import dirname
from os.path import splitext
from os.path import abspath

from PyQt5.QtCore import Qt
from PyQt5.QtCore import QPoint
from PyQt5.QtCore import QMimeData
from PyQt5.QtGui import QDrag
from PyQt5.QtGui import QImage
from PyQt5.QtGui import QPixmap
from PyQt5.QtGui import QPainter
from PyQt5.QtGui import QKeySequence
from PyQt5.QtGui import QImageReader
from PyQt5.QtWidgets import QMenu
from PyQt5.QtWidgets import QWidget
from PyQt5.QtWidgets import QLabel
from PyQt5.QtWidgets import QPushButton
from PyQt5.QtWidgets import QComboBox
from PyQt5.QtWidgets import QSpinBox
from PyQt5.QtWidgets import QActionGroup
from PyQt5.QtWidgets import QSizePolicy
from PyQt5.QtWidgets import QTableWidget
from PyQt5.QtWidgets import QHBoxLayout
from PyQt5.QtWidgets import QVBoxLayout
from PyQt5.QtWidgets import QFileDialog
from PyQt5.QtWidgets import QApplication
from PyQt5.QtPrintSupport import QPrinter

from docx import Document
from docx.shared import Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

from Sisyphe.core.sisypheConstants import getBitmapExt
from Sisyphe.widgets.basicWidgets import messageBox
from Sisyphe.widgets.basicWidgets import LabeledComboBox

"""
Class hierarchy
~~~~~~~~~~~~~~~
    
    - QWidget -> ScreenshotsGridWidget
"""


class ScreenshotsGridWidget(QWidget):
    """
    ScreenshotsGridWidget

    Description
    ~~~~~~~~~~~

    Screenshots management.

    Inheritance
    ~~~~~~~~~~~

    QWidget -> ScreenshotsGridWidget

    Creation: 10/05/2022
    Last revision: 25/11/2024
    """

    _ROWS = 5
    _COLS = 5
    _PORTRAIT = 0
    _LANDSCAPE = 1
    _BITMAP = ('bmp', 'gif', 'jpeg', 'jpg', 'pbm', 'pgm',
               'png', 'ppm', 'svg', 'svgz', 'tga', 'tif',
               'tiff', 'wbmp', 'webp', 'xbm', 'xpm')

    # Special methods

    """
    Private attribute

    _rows           int, row count
    _columns        int, column count
    _selection      list[int, int], row, column selection
    _pixmaps        list[QPixmap | None], list of QPixmap displayed in table
    _orientation    int, 0 = portrait or 1 = landscape
    _table          QTableWidget
    _popup          QMenu
    _dragpos0       list[int, int], starting drag event, first mouse position
    """

    def __init__(self, parent=None):
        super().__init__(parent)

        self._rows = self._ROWS
        self._columns = self._COLS
        self._selection = [0, 0]
        self._pixmaps: list[QPixmap | None] = list()
        self._orientation = self._PORTRAIT
        self._dragpos0 = None
        self._background = 0

        # Init QTableWidget

        self._table = QTableWidget()
        self._table.setRowCount(self._ROWS)
        self._table.setColumnCount(self._COLS)
        self._table.verticalHeader().setVisible(False)
        self._table.horizontalHeader().setVisible(False)
        self._table.horizontalScrollBar().setVisible(False)
        self._table.verticalScrollBar().setVisible(False)
        self._table.setShowGrid(False)
        # < Revision 26/04/2025
        # self._table.setEnabled(False)
        self._table.setAttribute(Qt.WA_TransparentForMouseEvents)
        # Revision 26/04/2025 >
        for i in range(0, self._ROWS):
            for j in range(0, self._COLS):
                self._pixmaps.append(None)
                w = QLabel()
                w.setFrameShape(QLabel.Box)
                w.setScaledContents(False)
                # noinspection PyTypeChecker
                w.setSizePolicy(QSizePolicy.Ignored, QSizePolicy.Ignored)
                w.setEnabled(True)
                if i == 0 and j == 0: w.setStyleSheet('background-color: black; border-width: 5px; border-color: rgb(0, 125, 255); margin: 0px;')
                else: w.setStyleSheet('background-color: black; border: none; border-color: rgb(0, 125, 255); margin: 0px;')
                self._table.setCellWidget(i, j, w)

        self._boxtable = QWidget()
        lyout = QVBoxLayout()
        lyout.setContentsMargins(0, 0, 0, 0)
        lyout.addWidget(self._table)
        self._boxtable.setLayout(lyout)

        # Init QMenu

        self._popup = QMenu()
        # noinspection PyTypeChecker
        self._popup.setWindowFlag(Qt.NoDropShadowWindowHint, True)
        # noinspection PyTypeChecker
        self._popup.setWindowFlag(Qt.FramelessWindowHint, True)
        self._popup.setAttribute(Qt.WA_TranslucentBackground, True)
        a = self._popup.addAction('Copy')
        # noinspection PyUnresolvedReferences
        a.triggered.connect(self.copyToClipboard)
        a = self._popup.addAction('Cut')
        # noinspection PyUnresolvedReferences
        a.triggered.connect(self.cutToClipboard)
        a = self._popup.addAction('Paste')
        # noinspection PyUnresolvedReferences
        a.triggered.connect(self.pasteFromClipboard)
        self._popup.addSeparator()
        a = self._popup.addAction('Add bitmap(s)...')
        # noinspection PyUnresolvedReferences
        a.triggered.connect(self.pasteFromDisk)
        a = self._popup.addAction('Clear selected')
        # noinspection PyUnresolvedReferences
        a.triggered.connect(self.deleteSelection)
        a = self._popup.addAction('Clear all')
        # noinspection PyUnresolvedReferences
        a.triggered.connect(self.clearAll)
        self._popup.addSeparator()
        submenu = self._popup.addMenu('Orientation')
        group = QActionGroup(self)
        group.setExclusive(True)
        a = submenu.addAction('Portrait')
        a.setCheckable(True)
        a.setActionGroup(group)
        a.setChecked(True)
        # noinspection PyUnresolvedReferences
        a.triggered.connect(self.setOrientationToPortrait)
        a = submenu.addAction('Landscape')
        a.setCheckable(True)
        a.setActionGroup(group)
        # noinspection PyUnresolvedReferences
        a.triggered.connect(self.setOrientationToLandscape)
        submenu = self._popup.addMenu('Row count')
        group = QActionGroup(self)
        group.setExclusive(True)
        for i in range(1, self._ROWS+1):
            a = submenu.addAction(str(i))
            a.setCheckable(True)
            a.setActionGroup(group)
            if i == 5: a.setChecked(True)
            # noinspection PyUnresolvedReferences
            a.triggered.connect(lambda dummy, v=i: self.setRows(v))
        submenu = self._popup.addMenu('Column count')
        group = QActionGroup(self)
        group.setExclusive(True)
        for i in range(1, self._COLS+1):
            a = submenu.addAction(str(i))
            a.setCheckable(True)
            a.setActionGroup(group)
            if i == 5: a.setChecked(True)
            # noinspection PyUnresolvedReferences
            a.triggered.connect(lambda dummy, v=i: self.setCols(v))
        self._popup.addSeparator()
        a = self._popup.addAction('Save bitmap...')
        # noinspection PyUnresolvedReferences
        a.triggered.connect(lambda dummy: self.saveToBitmap())
        a = self._popup.addAction('Save...')
        # noinspection PyUnresolvedReferences
        a.triggered.connect(lambda dummy: self.save())

        # Init buttons

        btlyout = QHBoxLayout()
        btlyout.setContentsMargins(0, 0, 0, 0)
        btlyout.setSpacing(5)
        btlyout.setAlignment(Qt.AlignCenter)

        w = QComboBox()
        w.setToolTip('Grid orientation, portrait or landscape.')
        w.addItem('Portrait')
        w.addItem('Landscape')
        w.setCurrentIndex(0)
        # noinspection PyUnresolvedReferences
        w.currentIndexChanged.connect(self.setOrientation)
        btlyout.addWidget(w)
        self._backg = LabeledComboBox(title='Background', fontsize=14)
        self._backg.addItem('black')
        self._backg.addItem('white')
        self._backg.setCurrentIndex(0)
        # noinspection PyUnresolvedReferences
        self._backg.currentIndexChanged.connect(self._backgroundChanged)
        btlyout.addWidget(self._backg)
        w = QLabel('Rows')
        btlyout.addWidget(w)
        w = QSpinBox()
        w.setToolTip('Rows count.')
        # < Revision 13/12/2024
        # w.setRange(0, 5)
        w.setRange(1, 5)
        # Revision 13/12/2024 >
        w.setValue(5)
        # noinspection PyUnresolvedReferences
        w.valueChanged.connect(self.setRows)
        btlyout.addWidget(w)
        w = QLabel('Columns')
        btlyout.addWidget(w)
        w = QSpinBox()
        w.setToolTip('Columns count.')
        # < Revision 13/12/2024
        # w.setRange(0, 5)
        w.setRange(1, 5)
        # Revision 13/12/2024 >
        w.setValue(5)
        # noinspection PyUnresolvedReferences
        w.valueChanged.connect(self.setCols)
        btlyout.addWidget(w)
        w = QPushButton('Add...')
        w.setToolTip('Add bitmap from disk.')
        # noinspection PyUnresolvedReferences
        w.clicked.connect(self.pasteFromDisk)
        btlyout.addWidget(w)
        w = QPushButton('Clear')
        w.setToolTip('Clear selected screenshot.')
        # noinspection PyUnresolvedReferences
        w.clicked.connect(self.deleteSelection)
        btlyout.addWidget(w)
        w = QPushButton('Clear all')
        w.setToolTip('Clear screenshots grid.')
        # noinspection PyUnresolvedReferences
        w.clicked.connect(self.clearAll)
        btlyout.addWidget(w)
        w = QPushButton('Save...')
        w.setToolTip('Save screenshots grid to word document, html or pdf.')
        # noinspection PyUnresolvedReferences
        w.clicked.connect(lambda dummy: self.save())
        btlyout.addWidget(w)
        w = QPushButton('Save bitmap...')
        w.setToolTip('Save screenshots grid to bitmap.')
        # noinspection PyUnresolvedReferences
        w.clicked.connect(lambda dummy: self.saveToBitmap())
        btlyout.addWidget(w)

        # Init layout

        lyout = QVBoxLayout()
        lyout.setContentsMargins(0, 0, 0, 0)
        lyout.setSpacing(0)
        lyout.addWidget(self._boxtable)
        lyout.addLayout(btlyout)
        self.setLayout(lyout)

        self.setOrientationToPortrait()
        self.setAcceptDrops(True)

    # Private methods

    def _updateSize(self):
        c = 297 / 210  # A4 ratio
        if self.isPortraitOrientation():
            h = self._boxtable.height()
            w = h / c
            if w > self._boxtable.width():
                w = self._boxtable.width()
                h = w * c
        else:
            w = self._boxtable.width()
            h = w / c
            if h > self._boxtable.height():
                h = self._boxtable.height()
                w = h * c
        w = int(w / self._columns)
        h = int(h / self._rows)
        # Update rows height
        for i in range(0, self._rows):
            self._table.setRowHeight(i, h)
        # Update columns width
        for i in range(0, self._columns):
            self._table.setColumnWidth(i, w)
        # Update table size and position
        h = h * self._rows
        if h > self._boxtable.height(): h = self._boxtable.height()
        w = w * self._columns
        if w > self._boxtable.width(): w = self._boxtable.width()
        self._table.resize(w, h)
        x = (self._boxtable.width() - w) // 2
        y = (self._boxtable.height() - h) // 2
        self._table.move(x, y)
        # Update pixmap size
        for i in range(0, self._rows):
            for j in range(0, self._columns):
                idx = self._getIndex(i, j)
                img = self._pixmaps[idx]
                if img is not None:
                    widget = self._table.cellWidget(i, j)
                    # noinspection PyTypeChecker
                    widget.setPixmap(img.scaled(widget.width(), widget.height(), Qt.KeepAspectRatio))

    def _updateCells(self, rows, cols):
        pixmaps = list()
        for i in range(0, self._rows):
            for j in range(0, self._columns):
                pixmaps.append(self._getPixmap(i, j))
        n = len(pixmaps)
        idx = 0
        for i in range(0, rows):
            for j in range(0, cols):
                if idx < n:
                    self._setPixmap(i, j, pixmaps[idx])
                    idx += 1
                else: self._setPixmap(i, j, None)
        self._rows = rows
        self._columns = cols
        self._updateSize()
        self.setSelection(0, 0)

    def _getPixmap(self, r, c):
        return self._pixmaps[self._getIndex(r, c)]

    def _getSelectedPixmap(self):
        return self._pixmaps[self._getSelectedIndex()]

    def _setPixmap(self, r, c, img):
        w = self._table.cellWidget(r, c)
        if img is not None:
            if isinstance(img, QPixmap):
                self._pixmaps[self._getIndex(r, c)] = img
                if img is not None: w.setPixmap(img.scaled(w.width(), w.height(), Qt.KeepAspectRatio))
            else: raise TypeError('image parameter type {} is not QPixmap.'.format(type(img)))
        else:
            w.clear()
            self._pixmaps[self._getIndex(r, c)] = None

    def _setPixmapToSelection(self, img):
        r, c = self.getSelection()
        self._setPixmap(r, c, img)

    def _nextCell(self):
        r, c = self._selection
        if c == self._columns - 1:
            if r == self._rows - 1: r, c = 0, 0
            else: c, r = 0, r + 1
        else: c += 1
        self.setSelection(r, c)

    def _previousCell(self):
        r, c = self._selection
        if c == 0:
            if r == 0: r, c = self._rows - 1, self._columns - 1
            else: c, r = self._columns - 1, r - 1
        else: c -= 1
        self.setSelection(r, c)

    def _belowCell(self):
        r, c = self._selection
        if r == self._rows - 1:
            if c == self._columns - 1: r, c = 0, 0
            else: c, r = c + 1, 0
        else: r += 1
        self.setSelection(r, c)

    def _aboveCell(self):
        r, c = self._selection
        if r == 0:
            if c == 0: r, c = self._rows - 1, self._columns - 1
            else: c, r = c - 1, self._rows - 1
        else: r -= 1
        self.setSelection(r, c)

    def _getIndex(self, r, c):
        return r * self._COLS + c

    def _getRowColumn(self, idx):
        r = idx // self._COLS
        c = idx - r * self._COLS
        return r, c

    def _getSelectedIndex(self):
        return self._getIndex(self._selection[0], self._selection[1])

    def _mapToTableViewport(self, p):
        p = self.mapToGlobal(p)
        return self._table.mapFromGlobal(p)

    def _backgroundChanged(self, index: int):
        if index == 0: self.setBlackBackground()
        else: self.setWhiteBackground()

    # Public methods

    # < Revision 13/12/2024
    def setWhiteBackground(self):
        self._background = 1
        si, sj = self.getSelection()
        for i in range(0, self._ROWS):
            for j in range(0, self._COLS):
                w = self._table.cellWidget(i, j)
                if i == si and j == sj: w.setStyleSheet('background-color: white; border-width: 5px; border-color: rgb(0, 125, 255); margin: 0px;')
                else: w.setStyleSheet('background-color: white; border: none; border-color: rgb(0, 125, 255); margin: 0px;')
    # Revision 13/12/2024 >

    # < Revision 13/12/2024
    def setBlackBackground(self):
        self._background = 0
        si, sj = self.getSelection()
        for i in range(0, self._ROWS):
            for j in range(0, self._COLS):
                w = self._table.cellWidget(i, j)
                if i == si and j == sj: w.setStyleSheet('background-color: black; border-width: 5px; border-color: rgb(0, 125, 255); margin: 0px;')
                else: w.setStyleSheet('background-color: black; border: none; border-color: rgb(0, 125, 255); margin: 0px;')
    # Revision 13/12/2024 >

    def getPopup(self):
        return self._popup

    def setRows(self, rows):
        if isinstance(rows, int):
            if rows < 6:
                for i in range(0, self._ROWS-1):
                    if i < rows: self._table.showRow(i)
                    else: self._table.hideRow(i)
                self._updateCells(rows, self._columns)
            else: raise ValueError('parameter value must be less than 5.')
        else: raise TypeError('parameter type {} is not int.'.format(type(rows)))

    def setCols(self, cols):
        if isinstance(cols, int):
            if cols < 6:
                for i in range(0, self._COLS-1):
                    if i < cols: self._table.showColumn(i)
                    else: self._table.hideColumn(i)
                self._updateCells(self._rows, cols)
            else: raise ValueError('parameter value must be less than 5.')
        else: raise TypeError('parameter type {} is not int.'.format(type(cols)))

    # noinspection PyUnusedLocal
    def setGridSize(self, rows, cols):
        if isinstance(rows, int):
            if rows < 6:
                for i in range(0, self._ROWS-1):
                    if i < rows: self._table.showRow(i)
                    else: self._table.hideRow(i)
            else: raise ValueError('parameter value must be less than 5.')
        else: raise TypeError('parameter type {} is not int.'.format(type(rows)))
        # noinspection PyUnreachableCode
        if isinstance(cols, int):
            if cols < 6:
                for i in range(0, self._COLS-1):
                    if i < cols: self._table.showColumn(i)
                    else: self._table.hideColumn(i)
            else: raise ValueError('parameter value must be less than 5.')
        else: raise TypeError('parameter type {} is not int.'.format(type(cols)))
        # noinspection PyUnreachableCode
        self._updateCells(rows, cols)

    def getRows(self):
        return self._rows

    def getCols(self):
        return self._cols

    def setOrientation(self, o):
        if isinstance(o, int):
            if o in [0, 1]:
                self._orientation = o
                self._updateSize()
            else: raise ValueError('parameter value {} is not in [0,1]'.format(o))
        else: raise TypeError('parameter type {} is not int.'.format(type(o)))

    def setOrientationToPortrait(self):
        self.setOrientation(self._PORTRAIT)

    def setOrientationToLandscape(self):
        self.setOrientation(self._LANDSCAPE)

    def getOrientation(self, string=False):
        if string:
            if self._orientation == self._PORTRAIT: return 'portrait'
            else: return 'landscape'
        else: return self._orientation

    def isPortraitOrientation(self):
        return self._orientation == self._PORTRAIT

    def isLandscapeOrientation(self):
        return self._orientation == self._LANDSCAPE

    def setSelection(self, r, c):
        if all(isinstance(i, int) for i in [r, c]):
            if 0 <= r < self._rows and 0 <= c < self._columns:
                rs, cs = self.getSelection()
                w = self._table.cellWidget(rs, cs)
                if self._background == 0: w.setStyleSheet('background-color: black; border: none; border-color: rgb(0, 125, 255); margin: 0px;')
                else: w.setStyleSheet('background-color: white; border: none; border-color: rgb(0, 125, 255); margin: 0px;')
                self._selection = [r, c]
                w = self._table.cellWidget(r, c)
                if self._background == 0: w.setStyleSheet('background-color: black; border-width: 5px; border-color: rgb(0, 125, 255); margin: 0px;')
                else: w.setStyleSheet('background-color: white; border-width: 5px; border-color: rgb(0, 125, 255); margin: 0px;')
            else: raise IndexError('Index is out of range.')
        else: raise TypeError('Parameters type is not int.')

    def getSelection(self):
        return self._selection

    def copyImageToCellFromIndex(self, row, col, img):
        if isinstance(img, str):
            if exists(img):
                ext = splitext(img)[1]
                if ext[1:] in QImageReader.supportedImageFormats():
                    img = QPixmap(img)
            else: raise ValueError('{} is not a bitmap file.'.format(basename(img)))
        else: raise FileNotFoundError('No such file {}.'.format(img))
        if isinstance(img, QImage): img = QPixmap.fromImage(img)
        if isinstance(img, QPixmap): self._setPixmap(row, col, img)
        else: raise TypeError('image parameter type {} is not str, QImage or QPixmap.'.format(type(img)))

    def copyImageToCellFromMouse(self, mousex, mousey, img):
        if isinstance(img, str):
            if exists(img):
                filt = ['.' + f.data().decode() for f in QImageReader.supportedImageFormats()]
                if splitext(img)[1] in filt:
                    img = QPixmap(img)
            else: raise ValueError('{} is not a bitmap file.'.format(basename(img)))
        else: raise FileNotFoundError('No such file {}.'.format(img))
        if isinstance(img, QPixmap):
            pt = self._table.PointToClient(QPoint(mousex, mousey))
            idx = self._table.indexAt(pt)
            self.copyImageToCellFromIndex(idx.row(), idx.column(), img)
        else: raise TypeError('image parameter type {} is not QPixmap.'.format(img))

    def copyToClipboard(self):
        clipboard = QApplication.clipboard()
        img = self._getSelectedPixmap()
        if img is not None: clipboard.setPixmap(img)

    def cutToClipboard(self):
        clipboard = QApplication.clipboard()
        img = self._getSelectedPixmap()
        if img is not None:
            clipboard.setPixmap(img)
            self.deleteSelection()

    def pasteFromClipboard(self):
        clipboard = QApplication.clipboard()
        img = clipboard.image()
        if img is not None: img = QPixmap.fromImage(img)
        else: img = clipboard.pixmap()
        if img is not None:
            self._setPixmapToSelection(img)
            self._nextCell()

    def pasteFromDisk(self):
        filt = ['*.' + f.data().decode() for f in QImageReader.supportedImageFormats()]
        filt = ' '.join(filt)
        filenames = QFileDialog.getOpenFileNames(self,
                                                 'Select bitmap image',
                                                 getcwd(),
                                                 'Bitmap ({})'.format(filt))[0]
        if len(filenames) > 0:
            chdir(dirname(filenames[0]))
            for filename in filenames:
                img = QPixmap(filename)
                self._setPixmapToSelection(img)
                self._nextCell()

    def paste(self, img):
        if isinstance(img, str):
            if exists(img):
                ext = splitext(img)[1]
                if ext[1:] in QImageReader.supportedImageFormats():
                    img = QPixmap(img)
            else: raise FileNotFoundError('No such file {}.'.format(img))
        if isinstance(img, QImage): img = QPixmap.fromImage(img)
        if isinstance(img, QPixmap):
            self._setPixmapToSelection(img)
            self._nextCell()
        else: raise TypeError('parameter type {} is not QPixmap or QImage'.format(type(img)))

    def delete(self, r, c):
        self._setPixmap(r, c, None)

    def deleteSelection(self):
        self._setPixmapToSelection(None)

    def clearAll(self):
        for i in range(0, len(self._pixmaps)):
            self._pixmaps[i] = None
            c = i // self._ROWS
            r = i - c * self._ROWS
            w = self._table.cellWidget(r, c)
            w.clear()

    def createHtml(self, filepath='', dpi=100):
        if isinstance(filepath, str):
            if not exists(filepath): filepath = getcwd()
            # A4 Portrait 8.3" x 11.7"
            # Calc image size
            wi = int(8.3 * dpi / self._columns)
            hi = int(11.7 * dpi / self._rows)
            if self.isLandscapeOrientation(): wi, hi = hi, wi
            # save temporary bitmap
            for i in range(0, len(self._pixmaps)):
                if self._pixmaps[i] is not None:
                    r, c = self._getRowColumn(i)
                    # noinspection PyTypeChecker
                    img = self._pixmaps[i].scaled(wi, hi, Qt.KeepAspectRatio)
                    filename = join(filepath, 'P{}{}.png'.format(r, c))
                    img.save(filename)
            # html str
            html = '<!DOCTYPE html>\n'
            html += '<html>\n'
            html += '\t<body>\n'
            if self._backg.currentIndex() == 0:
                html += '\t\t<table style=\"background-color: black; border: none; empty-cell: show;\">\n'
            else: html += '\t\t<table style=\"background-color: white; border: none; empty-cell: show;\">\n'
            for i in range(0, self._rows):
                html += '\t\t\t<tr>\n'
                for j in range(0, self._columns):
                    filename = join(filepath, 'P{}{}.png'.format(i, j))
                    html += '\t\t\t\t<td>\n'
                    if exists(filename):
                        style = 'border: none; width: {}px; height: {}px;'.format(wi, hi)
                        html += '\t\t\t\t\t<img src=\"{}\" alt=\"\" style=\"{}\"></img>\n'.format(filename, style)
                    html += '\t\t\t\t</td>\n'
                html += '\t\t\t</tr>\n'
            html += '\t\t</table>\n'
            html += '\t</body>\n'
            html += '</html>\n'
            return html
        else: raise TypeError('parameter type {} is not str.'.format(type(filepath)))

    def saveToHtml(self, filename=None, dpi=100):
        if filename is None:
            filename = QFileDialog.getSaveFileName(self, 'Save screenshots to Html', getcwd(), 'Html (*.html)')
            filename = filename[0]
        if filename != '':
            # noinspection PyTypeChecker
            ext = splitext(filename)[1]
            if ext != '.html':
                # noinspection PyTypeChecker
                filename = join(basename(filename), '.html')
            html = self.createHtml(dirname(filename), dpi)
            f = open(filename, "w+")
            try: f.write(html)
            except Exception as err:
                messageBox(self, 'Save screenshot to html.', text='error : {}'.format(err))
            finally: f.close()

    def saveToPdf(self,  filename=None, dpi=100):
        if filename is None:
            filename = QFileDialog.getSaveFileName(self, 'Save screenshots to Pdf', getcwd(), 'Pdf (*.pdf)')
            filename = filename[0]
        if filename != '':
            # noinspection PyTypeChecker
            ext = splitext(filename)[1]
            if ext != '.pdf':
                # noinspection PyTypeChecker
                filename = join(basename(filename), '.pdf')
            try:
                printer = QPrinter()
                if self._orientation == 0: orient = QPrinter.Portrait
                else: orient = QPrinter.Landscape
                printer.setOrientation(orient)
                printer.setResolution(dpi)
                printer.setPageSize(QPrinter.A4)
                printer.setOutputFormat(QPrinter.PdfFormat)
                printer.setOutputFileName(filename)
                painter = QPainter(printer)
                # Page size, with 50 margin
                wp = painter.viewport().width() - 100
                hp = painter.viewport().height() - 100
                # Calc image size
                wi = wp // self._columns
                hi = hp // self._rows
                # Copy images
                if self._backg.currentIndex() == 0:
                    painter.fillRect(50, 50, wp, hp, Qt.black)
                for i in range(0, self._rows):
                    for j in range(0, self._columns):
                        idx = self._getIndex(i, j)
                        x0 = 50 + j * wi
                        y0 = 50 + i * hi
                        if self._backg.currentIndex() == 0:
                            painter.fillRect(x0, y0, wi, hi, Qt.black)
                        if self._pixmaps[idx] is not None:
                            # noinspection PyTypeChecker
                            img = self._pixmaps[idx].scaled(wi, hi, Qt.KeepAspectRatio)
                            x = x0 + int((wi - img.width()) / 2)
                            y = y0 + int((hi - img.height()) / 2)
                            painter.drawPixmap(x, y, img)
                painter.end()
            except Exception as err:
                messageBox(self, 'Save screenshots to pdf.', text='error : {}'.format(err))

    def saveToDocx(self,  filename=None, dpi=60):
        if filename is None:
            filename = QFileDialog.getSaveFileName(self, 'Save screenshots to Word document',
                                                   getcwd(), 'Word (*.docx)')
            filename = filename[0]
        if filename != '':
            # noinspection PyTypeChecker
            chdir(dirname(filename))
            # noinspection PyTypeChecker
            tempfile = join(dirname(filename), 'temp.png')
            self.saveToBitmap(tempfile, dpi)
            # Generate docx
            docx = Document()
            sections = docx.sections
            for section in sections:
                section.top_margin = Cm(1.0)
                section.bottom_margin = Cm(1.0)
                section.left_margin = Cm(1.0)
                section.right_margin = Cm(1.0)
            docx.add_picture(tempfile)
            docx.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            try:
                # noinspection PyTypeChecker
                docx.save(filename)
            except Exception as err:
                messageBox(self, 'Save screenshots to word document.', text='error : {}'.format(err))
            # Delete temporary bitmap
            if exists(tempfile): remove(tempfile)

    def saveToBitmap(self, filename=None, dpi=100):
        # Calc image size
        wi = int(8.3 * dpi / self._columns)
        hi = int(11.7 * dpi / self._rows)
        if self.isLandscapeOrientation(): wi, hi = hi, wi
        wp = wi * self._columns
        hp = hi * self._rows
        # Copy images
        page = QPixmap(wp, hp)
        painter = QPainter(page)
        if self._backg.currentIndex() == 0:
            painter.fillRect(0, 0, wp, hp, Qt.black)
        for i in range(0, self._rows):
            for j in range(0, self._columns):
                idx = self._getIndex(i, j)
                if self._pixmaps[idx] is not None:
                    # noinspection PyTypeChecker
                    img = self._pixmaps[idx].scaled(wi, hi, Qt.KeepAspectRatio)
                    x = j * wi + int((wi - img.width()) / 2)
                    y = i * hi + int((hi - img.height()) / 2)
                    painter.drawPixmap(x, y, img)
        painter.end()
        # Save image
        if filename is None:
            filt = 'BMP (*.bmp);;GIF (*.gif);;JPEG (*.jpeg);;JPG (*.jpg);;' \
                   'PNG (*.png);;TIF (*.tif);;TIFF (*.tiff);;WEBP (*.webp)'
            filename = QFileDialog.getSaveFileName(self, 'Save screenshots to bitmap', getcwd(),
                                                   filt, 'PNG (*.png)')
            filename = filename[0]
        if filename != '':
            # noinspection PyTypeChecker
            chdir(dirname(filename))
            # noinspection PyTypeChecker
            ext = splitext(filename)[1]
            if ext in getBitmapExt():
                try: page.save(filename)
                except Exception as err:
                    messageBox(self, 'Save screenshots to bitmap', text='error : {}'.format(err))

    def save(self, filename=None, dpi=100):
        if filename is None:
            filename = QFileDialog.getSaveFileName(self, 'Save screenshots', getcwd(),
                                                   'Html (*.html);;Pdf (*.pdf);;Word (*.docx)', 'Pdf (*.pdf)')
            filename = filename[0]
        if filename != '':
            # noinspection PyTypeChecker
            chdir(dirname(filename))
            # noinspection PyTypeChecker
            ext = splitext(filename)[1]
            filt = ['.html', '.pdf', '.docx']
            if ext in filt:
                if ext == filt[0]: self.saveToHtml(filename, dpi)
                elif ext == filt[1]: self.saveToPdf(filename, dpi)
                elif ext == filt[2]: self.saveToDocx(filename)

    # Qt events

    def paintEvent(self, event):
        self._updateSize()
        super().paintEvent(event)

    def resizeEvent(self, event):
        super().resizeEvent(event)
        self._updateSize()

    def keyPressEvent(self, event):
        if event.matches(QKeySequence.Copy): self.copyToClipboard()
        elif event.matches(QKeySequence.Cut): self.cutToClipboard()
        elif event.matches(QKeySequence.Paste): self.pasteFromClipboard()
        elif event.key() == Qt.Key_Home: self.setSelection(0, 0)
        elif event.key() == Qt.Key_End: self.setSelection(self._rows, self._columns)
        elif event.key() in [Qt.Key_Tab, Qt.Key_Right]: self._nextCell()
        elif event.key() == Qt.Key_Left: self._previousCell()
        elif event.key() == Qt.Key_Up: self._aboveCell()
        elif event.key() == Qt.Key_Down: self._belowCell()
        elif event.key() == Qt.Key_Delete: self.deleteSelection()
        elif event.key() == Qt.Key_Backspace:
            self.deleteSelection()
            self._previousCell()
        super().keyPressEvent(event)

    def dragEnterEvent(self, event):
        if event.mimeData().hasText() or event.mimeData().hasImage(): event.accept()
        else: event.ignore()

    def dropEvent(self, event):
        img = None
        if event.mimeData().hasText():
            txt = event.mimeData().text()
            filename = abspath('/'.join([v for v in txt.split('/') if v not in ('file:', '')]))
            if exists(filename):
                ext = splitext(filename)[1]
                if ext[1:] in self._BITMAP: img = QPixmap(filename)
        elif event.mimeData().hasImage():
            r, c = self.getSelection()
            idx = self._table.indexAt(self._mapToTableViewport(event.pos()))
            if idx.row() != r or idx.column() != c:
                if 0 <= idx.row() < self._rows and 0 <= idx.column() < self._columns:
                    self.delete(r, c)
                    img = event.mimeData().imageData()
        if img is not None:
            idx = self._table.indexAt(self._mapToTableViewport(event.pos()))
            if 0 <= idx.row() < self._rows and 0 <= idx.column() < self._columns:
                self._setPixmap(idx.row(), idx.column(), img)
                self.setSelection(idx.row(), idx.column())
                event.accept()
            else: self._setPixmapToSelection(img)

    def contextMenuEvent(self, event):
        # Popup
        self._popup.exec(event.globalPos())

    def mousePressEvent(self, event):
        idx = self._table.indexAt(self._mapToTableViewport(event.pos()))
        # Selection
        if 0 <= idx.row() < self._rows and 0 <= idx.column() < self._columns:
            self.setSelection(idx.row(), idx.column())
        # Starting Drag
        if event.button() == Qt.LeftButton:
            self._dragpos0 = event.pos()

    def mouseReleaseEvent(self, event):
        self._dragpos0 = None

    def mouseMoveEvent(self, event):
        if self._dragpos0 is not None:
            if (event.pos() - self._dragpos0).manhattanLength() > 20:
                idx = self._table.indexAt(self._mapToTableViewport(self._dragpos0))
                img = self._getPixmap(idx.row(), idx.column())
                if img is not None:
                    data = QMimeData()
                    data.setImageData(img)
                    drag = QDrag(self)
                    drag.setMimeData(data)
                    w = self._table.cellWidget(idx.row(), idx.column())
                    icn = w.pixmap()
                    drag.setPixmap(icn)
                    drag.setHotSpot(QPoint(icn.width() // 2, icn.height() // 2))
                    drag.exec(Qt.MoveAction)
                    self._dragpos0 = None
